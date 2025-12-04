
import os
import json
import pandas as pd
from typing import List, Dict, Any
from dotenv import load_dotenv
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field

# Load environment variables
load_dotenv()

class ESGMetrics(BaseModel):
    company_name: str = Field(description="Name of the company")
    scope1_emissions: float = Field(description="Scope 1 GHG emissions (tCO2eq)", default=0.0)
    scope2_emissions: float = Field(description="Scope 2 GHG emissions (tCO2eq)", default=0.0)
    energy_consumption: float = Field(description="Total energy consumption (TJ or MWh converted to TJ)", default=0.0)
    safety_accident_rate: float = Field(description="Safety accident rate (LTIFR or similar)", default=0.0)
    training_hours_per_employee: float = Field(description="Average training hours per employee", default=0.0)
    board_independence_ratio: float = Field(description="Ratio of independent directors (%)", default=0.0)

class DataExtractor:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-m3",
            model_kwargs={'device': 'cpu'}
        )
        self.vectorstore = Chroma(
            persist_directory="./chroma_db",
            embedding_function=self.embeddings
        )
        self.llm = ChatOpenAI(
            model="gpt-4o",
            temperature=0,
            api_key=os.getenv("OPENAI_API_KEY")
        )
        self.parser = JsonOutputParser(pydantic_object=ESGMetrics)

    def get_all_companies(self) -> List[str]:
        """Get list of unique companies from ChromaDB metadata"""
        data = self.vectorstore.get(include=['metadatas'])
        companies = set()
        for meta in data['metadatas']:
            if meta and 'company' in meta:
                companies.add(meta['company'])
        return list(companies)

    def extract_company_metrics(self, company_name: str) -> Dict[str, Any]:
        """Extract ESG metrics for a specific company using RAG + LLM"""
        print(f"  - Extracting metrics for: {company_name}")
        
        # 1. Retrieve relevant documents
        # We search for a broad query to get a mix of E/S/G data
        query = f"{company_name} 온실가스 배출량 Scope 1 2 에너지 사용량 안전사고율 교육시간 이사회 독립성 비율"
        docs = self.vectorstore.similarity_search(query, k=10, filter={"company": company_name})
        
        context = "\n\n".join([doc.page_content for doc in docs])
        
        # 2. Extract using LLM
        prompt = PromptTemplate(
            template="""You are an expert data analyst. Extract the following ESG metrics for the company '{company_name}' from the provided context.
            
            Context:
            {context}
            
            Instructions:
            - Extract numerical values only. Remove units (e.g., '1,234 tCO2eq' -> 1234).
            - If a value is not found, return 0.0.
            - For 'board_independence_ratio', calculate it if not explicitly stated (Independent Directors / Total Directors * 100).
            
            {format_instructions}
            """,
            input_variables=["company_name", "context"],
            partial_variables={"format_instructions": self.parser.get_format_instructions()}
        )
        
        chain = prompt | self.llm | self.parser
        
        try:
            metrics = chain.invoke({"company_name": company_name, "context": context})
            return metrics
        except Exception as e:
            print(f"    Error extracting metrics for {company_name}: {e}")
            return ESGMetrics(company_name=company_name).dict()

    def run(self) -> pd.DataFrame:
        print("[Step 1] Data Collection Started...")
        companies = self.get_all_companies()
        print(f"  - Found {len(companies)} companies: {companies}")
        
        all_metrics = []
        for company in companies:
            metrics = self.extract_company_metrics(company)
            all_metrics.append(metrics)
            
        df = pd.DataFrame(all_metrics)
        print("[Step 1] Data Collection Completed.")
        return df


class StatisticalAnalyzer:
    def analyze(self, df: pd.DataFrame) -> Dict[str, Any]:
        print("[Step 2] Statistical Analysis Started...")
        analysis = {}
        
        # Calculate averages and medians
        numeric_cols = df.select_dtypes(include=['float64', 'int']).columns
        for col in numeric_cols:
            analysis[f"{col}_avg"] = df[col].mean()
            analysis[f"{col}_median"] = df[col].median()
            analysis[f"{col}_max"] = df[col].max()
            
            # Identify Best Practice (Top performer)
            # For emissions/accidents, lower is better. For others, higher is better.
            if "emissions" in col or "accident" in col or "consumption" in col:
                best_company = df.loc[df[col].idxmin()]['company_name']
                best_val = df[col].min()
            else:
                best_company = df.loc[df[col].idxmax()]['company_name']
                best_val = df[col].max()
                
            analysis[f"{col}_best_company"] = best_company
            analysis[f"{col}_best_value"] = best_val
            
        print("[Step 2] Analysis Completed.")
        return analysis

class ContentGenerator:
    def __init__(self):
        self.llm = ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=os.getenv("OPENAI_API_KEY"))

    def generate_executive_summary(self, analysis: Dict[str, Any]) -> str:
        prompt = f"""
        Write an Executive Summary for the '2024 Construction Industry Standard ESG Report'.
        
        Industry Statistics:
        - Average Scope 1 Emissions: {analysis.get('scope1_emissions_avg', 0):.1f} tCO2eq
        - Average Safety Accident Rate: {analysis.get('safety_accident_rate_avg', 0):.2f}
        - Best Safety Performer: {analysis.get('safety_accident_rate_best_company', 'N/A')}
        
        Tone: Professional, Objective, Data-driven (SustainBest style).
        Language: Korean.
        Length: 300 words.
        """
        return self.llm.invoke(prompt).content

    def generate_section(self, section_name: str, analysis: Dict[str, Any], best_company: str) -> str:
        prompt = f"""
        Write the '{section_name}' section for the ESG report.
        
        Context:
        - Industry Average: {analysis.get(f'{section_name}_avg', 'N/A')}
        - Best Practice Company: {best_company}
        
        Structure:
        1. Industry Status (Current state based on data)
        2. Best Practice Case Study (Focus on {best_company}'s strategy)
        3. Key Improvement Tasks for the Industry
        
        Language: Korean.
        """
        return self.llm.invoke(prompt).content

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentBuilder:
    def build(self, content: Dict[str, str], filename: str = "건설업계_표준_ESG보고서_2024.docx"):
        print("[Step 4] Document Assembly Started...")
        doc = Document()
        
        # Title
        title = doc.add_heading('2024 건설업계 표준 ESG 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(content['executive_summary'])
        
        # 2. Environmental
        doc.add_heading('2. Environmental (환경)', level=1)
        doc.add_paragraph(content['env_section'])
        
        # 3. Social
        doc.add_heading('3. Social (사회)', level=1)
        doc.add_paragraph(content['social_section'])
        
        # 4. Governance
        doc.add_heading('4. Governance (지배구조)', level=1)
        doc.add_paragraph(content['gov_section'])
        
        # Save
        doc.save(filename)
        print(f"[Step 4] Document Saved: {filename}")

def main():
    # Step 1: Data Collection
    extractor = DataExtractor()
    df = extractor.run()
    
    # Step 2: Analysis
    analyzer = StatisticalAnalyzer()
    analysis = analyzer.analyze(df)
    
    # Step 3: Content Generation
    generator = ContentGenerator()
    content = {
        'executive_summary': generator.generate_executive_summary(analysis),
        'env_section': generator.generate_section('scope1_emissions', analysis, analysis.get('scope1_emissions_best_company')),
        'social_section': generator.generate_section('safety_accident_rate', analysis, analysis.get('safety_accident_rate_best_company')),
        'gov_section': generator.generate_section('board_independence_ratio', analysis, analysis.get('board_independence_ratio_best_company'))
    }
    
    # Step 4: Document Assembly
    builder = DocumentBuilder()
    builder.build(content)

if __name__ == "__main__":
    main()

